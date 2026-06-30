import io
import fitz
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from presidio_analyzer import AnalyzerEngine, RecognizerResult
from presidio_anonymizer import AnonymizerEngine

analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()
CONFIDENCE_THRESHOLD = 0.5

def get_entity_priority(entity: RecognizerResult) -> int:
    """
    Return a priority for entity types – higher means preferred when overlapping.
    """
    priority = {
        "EMAIL_ADDRESS": 3,
        "PHONE_NUMBER": 2,
        "PERSON": 2,
        "LOCATION": 2,
        "URL": 1,
    }
    return priority.get(entity.entity_type, 0)

def merge_overlapping_entities(results: list) -> list:
    """
    Merge overlapping or contained entities, keeping the most relevant one.
    Priority: higher confidence > higher entity priority > longer span.
    """
    if not results:
        return []
    # Sort by start, then by end (shorter first)
    sorted_results = sorted(results, key=lambda x: (x.start, x.end))
    merged = []
    for entity in sorted_results:
        if not merged:
            merged.append(entity)
            continue
        last = merged[-1]
        if entity.start < last.end:
            # Overlap or containment – decide which to keep
            if entity.score > last.score:
                merged[-1] = entity
            elif entity.score == last.score:
                if get_entity_priority(entity) > get_entity_priority(last):
                    merged[-1] = entity
                elif get_entity_priority(entity) == get_entity_priority(last):
                    if (entity.end - entity.start) > (last.end - last.start):
                        merged[-1] = entity
        else:
            merged.append(entity)
    return merged

def redact_pdf(pdf_bytes: bytes, overrides: dict = None) -> bytes:
    """
    Redact PII from PDF, applying overrides if provided.
    overrides: dict mapping text_segment -> "REDACTED" or "KEPT_VISIBLE"
    """
    overrides = overrides or {}
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # Extract full text for detection
    full_text = ""
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text")
        full_text += page_text + "\n"

    if not full_text.strip():
        doc.close()
        raise ValueError("PDF contains no selectable text.")

    # Detect entities and merge overlapping
    raw_results = analyzer.analyze(text=full_text, language="en")
    results = merge_overlapping_entities(raw_results)

    # 1. Redact detected entities unless overridden to KEPT_VISIBLE
    for entity in results:
        entity_text = full_text[entity.start:entity.end]
        # Check if overridden
        if entity_text in overrides and overrides[entity_text] == "KEPT_VISIBLE":
            continue  # skip redaction
        # Redact all occurrences on all pages
        for page_num, page in enumerate(doc):
            quads = page.search_for(entity_text, quads=True)
            if quads:
                for quad in quads:
                    rect = quad.rect
                    page.add_redact_annot(rect, fill=(0, 0, 0))

    # 2. Apply forced redactions from overrides (entities not detected)
    for text_segment, action in overrides.items():
        if action == "REDACTED":
            # Try to redact this text everywhere it appears
            for page_num, page in enumerate(doc):
                quads = page.search_for(text_segment, quads=True)
                if quads:
                    for quad in quads:
                        rect = quad.rect
                        page.add_redact_annot(rect, fill=(0, 0, 0))

    # Apply redactions
    for page in doc:
        page.apply_redactions()

    output = io.BytesIO()
    doc.save(output, garbage=4, deflate=True)
    doc.close()
    return output.getvalue()


def redact_docx(docx_bytes: bytes, overrides: dict = None) -> bytes:
    """
    Redact PII from DOCX, applying overrides.
    overrides: dict mapping text_segment -> "REDACTED" or "KEPT_VISIBLE"
    """
    overrides = overrides or {}
    doc = Document(io.BytesIO(docx_bytes))

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if not full_text.strip():
        return docx_bytes

    # Detect entities and merge overlapping
    raw_results = analyzer.analyze(text=full_text, language="en")
    results = merge_overlapping_entities(raw_results)

    # Build a set of text segments that should be redacted
    redact_set = set()
    keep_visible_set = set()

    for entity in results:
        entity_text = full_text[entity.start:entity.end]
        if entity_text in overrides and overrides[entity_text] == "KEPT_VISIBLE":
            keep_visible_set.add(entity_text)
        else:
            redact_set.add(entity_text)

    # Also add any forced redactions from overrides (text not detected)
    for text_segment, action in overrides.items():
        if action == "REDACTED":
            redact_set.add(text_segment)

    # Now apply redactions to the document
    for paragraph in doc.paragraphs:
        para_text = paragraph.text
        if not para_text:
            continue
        # Check if any target is in this paragraph
        needs_redaction = any(target in para_text for target in redact_set)
        if not needs_redaction:
            continue

        # Preserve formatting from first run if exists
        if paragraph.runs:
            template_run = paragraph.runs[0]
            new_text = para_text
            for target in redact_set:
                new_text = new_text.replace(target, "[REDACTED]")
            paragraph.clear()
            new_run = paragraph.add_run(new_text)
            # Copy formatting
            new_run.bold = template_run.bold
            new_run.italic = template_run.italic
            new_run.underline = template_run.underline
            # Apply white text with black highlight
            new_run.font.color.rgb = RGBColor(255, 255, 255)
            new_run.font.highlight_color = WD_COLOR_INDEX.BLACK
        else:
            # No runs: set text directly
            for target in redact_set:
                if target in paragraph.text:
                    paragraph.text = paragraph.text.replace(target, "[REDACTED]")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()